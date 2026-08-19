#!/usr/bin/env python3
"""Build an editable Korean ebook DOCX from book.json."""
import json, sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

def rgb(value): return RGBColor.from_string(value.lstrip('#').upper())
def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill.lstrip('#').upper())
def set_cell_margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',180),('start',220),('bottom',180),('end',220)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')
def add_field(paragraph, instruction):
    run=paragraph.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=instruction
    sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate'); txt=OxmlElement('w:t'); txt.text='Word에서 필드를 업데이트하세요.'; sep.append(txt)
    end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    for n in (begin,instr,sep,end): run._r.append(n)
def add_image(doc,path,base,width_cm=12.2):
    if not path: return False
    p=Path(path); p=p if p.is_absolute() else base/p
    if not p.exists(): return False
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER; para.add_run().add_picture(str(p),width=Cm(width_cm)); return True
def add_block(doc,b,c):
    t=b.get('type','paragraph')
    if t=='paragraph': doc.add_paragraph(b.get('text',''),'Normal')
    elif t=='heading': doc.add_heading(b.get('text',''),level=int(b.get('level',2)))
    elif t=='quote':
        p=doc.add_paragraph(b.get('text',''),'Quote'); p.paragraph_format.left_indent=Cm(.7)
    elif t=='bullets':
        for x in b.get('items',[]): doc.add_paragraph(x,'List Bullet')
    elif t=='checklist':
        if b.get('title'): doc.add_heading(b['title'],level=3)
        for x in b.get('items',[]): doc.add_paragraph('☐ '+x,'List Paragraph')
    elif t=='callout':
        table=doc.add_table(rows=1,cols=1); cell=table.cell(0,0); shade(cell,c['muted']); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]
        if b.get('title'): r=p.add_run(b['title']+'\n'); r.bold=True; r.font.color.rgb=rgb(c['accent'])
        p.add_run(b.get('text','')); doc.add_paragraph()
    elif t=='page_break': doc.add_page_break()
def setup(doc,d):
    sec=doc.sections[0]; sec.page_width=Cm(14.8); sec.page_height=Cm(21); sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
    font=d.get('font','Malgun Gothic'); accent=d.get('accent','#7C5AA6'); text=d.get('text','#2C2630')
    for name,size,bold,color in [('Normal',10.5,False,text),('Title',30,True,accent),('Subtitle',13,False,text),('Heading 1',21,True,accent),('Heading 2',15,True,text),('Heading 3',12,True,accent),('Quote',11,False,accent)]:
        s=doc.styles[name]; s.font.name=font; s._element.rPr.rFonts.set(qn('w:eastAsia'),font); s.font.size=Pt(size); s.font.bold=bold; s.font.color.rgb=rgb(color)
    doc.styles['Normal'].paragraph_format.space_after=Pt(7); doc.styles['Normal'].paragraph_format.line_spacing=1.45
    for n in ('Heading 1','Heading 2','Heading 3'): doc.styles[n].paragraph_format.keep_with_next=True; doc.styles[n].paragraph_format.space_before=Pt(14); doc.styles[n].paragraph_format.space_after=Pt(8)
    return {'accent':accent,'text':text,'muted':d.get('muted','#F2ECF6')}
def main():
    if len(sys.argv)!=3: raise SystemExit('Usage: build_ebook.py book.json output.docx')
    spec=Path(sys.argv[1]).resolve(); out=Path(sys.argv[2]).resolve(); data=json.loads(spec.read_text(encoding='utf-8')); base=spec.parent; meta=data.get('meta',{}); design=data.get('design',{}); doc=Document(); colors=setup(doc,design)
    add_image(doc,design.get('cover_image'),base); p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(meta.get('title','제목'))
    if meta.get('subtitle'): p=doc.add_paragraph(meta['subtitle'],style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(meta.get('author',''),style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; doc.add_page_break()
    doc.add_heading('목차',level=1); add_field(doc.add_paragraph(),'TOC \\o "1-3" \\h \\z \\u'); doc.add_page_break(); add_image(doc,design.get('copyright_background'),base); doc.add_heading('판권',level=1)
    for label,value in [('제목',meta.get('title')),('저자',meta.get('author')),('브랜드',meta.get('brand')),('발행처',meta.get('publisher')),('발행일',meta.get('publication_date')),('ISBN',meta.get('isbn')),('연락처',meta.get('contact'))]:
        if value: doc.add_paragraph(f'{label}  {value}')
    if meta.get('copyright'): doc.add_paragraph(meta['copyright'])
    if meta.get('disclaimer'): doc.add_paragraph(meta['disclaimer'])
    doc.add_page_break()
    for b in data.get('front_matter',[]): add_block(doc,b,colors)
    if data.get('front_matter'): doc.add_page_break()
    for i,ch in enumerate(data.get('chapters',[]),1):
        if i>1: doc.add_page_break()
        add_image(doc,ch.get('image'),base); doc.add_heading(f'{ch.get("number",i)}장  {ch.get("title","")}',level=1)
        if ch.get('subtitle'): p=doc.add_paragraph(ch['subtitle'],style='Subtitle'); p.paragraph_format.keep_with_next=True
        for b in ch.get('blocks',[]): add_block(doc,b,colors)
    if data.get('back_matter'): doc.add_page_break()
    for b in data.get('back_matter',[]): add_block(doc,b,colors)
    footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_field(footer,'PAGE')
    out.parent.mkdir(parents=True,exist_ok=True); doc.save(out); print(out)
if __name__=='__main__': main()
